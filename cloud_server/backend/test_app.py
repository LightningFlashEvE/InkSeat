import base64
import copy
import hashlib
import io
import json
import os
import sys
import tempfile
import types
import unittest
import zipfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from PIL import Image


BACKEND_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(BACKEND_DIR))
os.environ['APP_SKIP_INIT_FOR_TESTS'] = '1'
for stream in (sys.stdout, sys.stderr):
    if hasattr(stream, 'reconfigure'):
        stream.reconfigure(encoding='utf-8', errors='replace')

try:
    import pymongo  # noqa: F401
except ModuleNotFoundError:
    pymongo_stub = types.ModuleType('pymongo')
    pymongo_errors_stub = types.ModuleType('pymongo.errors')

    class DuplicateKeyError(Exception):
        pass

    class MongoClient:
        def __init__(self, *args, **kwargs):
            raise AssertionError('Unit tests must not connect to MongoDB')

    pymongo_stub.MongoClient = MongoClient
    pymongo_stub.ReturnDocument = SimpleNamespace(AFTER=True)
    pymongo_errors_stub.DuplicateKeyError = DuplicateKeyError
    pymongo_stub.errors = pymongo_errors_stub
    sys.modules['pymongo'] = pymongo_stub
    sys.modules['pymongo.errors'] = pymongo_errors_stub

import app as backend  # noqa: E402
import config as backend_config  # noqa: E402
import db_indexes  # noqa: E402
import template_renderer  # noqa: E402


def make_test_logo_data_url(color=(0, 0, 255, 255), size=(40, 20)):
    image = Image.new('RGBA', size, color)
    buffer = io.BytesIO()
    image.save(buffer, format='PNG')
    encoded = base64.b64encode(buffer.getvalue()).decode('ascii')
    return f'data:image/png;base64,{encoded}'


def serialize_test_docx(document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class FakeCursor:
    def __init__(self, documents):
        self.documents = copy.deepcopy(documents)

    def sort(self, key, direction=1):
        self.documents.sort(
            key=lambda document: (document.get(key) is not None, document.get(key)),
            reverse=direction < 0,
        )
        return self

    def limit(self, count):
        self.documents = self.documents[:count]
        return self

    def __iter__(self):
        return iter(copy.deepcopy(self.documents))


class FakeCollection:
    def __init__(self, documents=None):
        self.documents = copy.deepcopy(documents or [])

    @staticmethod
    def _matches(document, query):
        for key, expected in (query or {}).items():
            if key == '$or':
                if not any(FakeCollection._matches(document, item) for item in expected):
                    return False
                continue
            actual = document.get(key)
            if isinstance(expected, dict):
                if '$exists' in expected and (key in document) != bool(expected['$exists']):
                    return False
                if '$gt' in expected and not (actual is not None and actual > expected['$gt']):
                    return False
                if '$gte' in expected and not (actual is not None and actual >= expected['$gte']):
                    return False
                if '$lte' in expected and not (actual is not None and actual <= expected['$lte']):
                    return False
                if '$ne' in expected and actual == expected['$ne']:
                    return False
                if '$in' in expected and actual not in expected['$in']:
                    return False
            elif actual != expected:
                return False
        return True

    @staticmethod
    def _project(document, projection):
        if not projection:
            return copy.deepcopy(document)
        included = [key for key, enabled in projection.items() if enabled and key != '_id']
        if included:
            result = {key: copy.deepcopy(document[key]) for key in included if key in document}
            if projection.get('_id', 1) and '_id' in document:
                result['_id'] = copy.deepcopy(document['_id'])
            return result
        result = copy.deepcopy(document)
        for key, enabled in projection.items():
            if not enabled:
                result.pop(key, None)
        return result

    def find_one(self, query=None, projection=None):
        for document in self.documents:
            if self._matches(document, query or {}):
                return self._project(document, projection)
        return None

    def find(self, query=None, projection=None):
        return FakeCursor([
            self._project(document, projection)
            for document in self.documents
            if self._matches(document, query or {})
        ])

    def insert_one(self, document):
        stored = copy.deepcopy(document)
        stored.setdefault('_id', len(self.documents) + 1)
        self.documents.append(stored)
        return SimpleNamespace(inserted_id=stored['_id'])

    def update_one(self, query, update, upsert=False):
        for index, document in enumerate(self.documents):
            if self._matches(document, query):
                self._apply_update(document, update, inserting=False)
                self.documents[index] = document
                return SimpleNamespace(matched_count=1, modified_count=1)
        if upsert:
            document = {key: value for key, value in query.items() if not key.startswith('$') and not isinstance(value, dict)}
            self._apply_update(document, update, inserting=True)
            self.insert_one(document)
            return SimpleNamespace(matched_count=0, modified_count=0, upserted_id=document.get('_id'))
        return SimpleNamespace(matched_count=0, modified_count=0)

    def find_one_and_delete(self, query):
        for index, document in enumerate(self.documents):
            if self._matches(document, query):
                return self.documents.pop(index)
        return None

    def find_one_and_update(self, query, update, upsert=False, return_document=None):
        for index, document in enumerate(self.documents):
            if self._matches(document, query):
                self._apply_update(document, update, inserting=False)
                self.documents[index] = document
                return copy.deepcopy(document)
        if upsert:
            document = {
                key: value
                for key, value in query.items()
                if not key.startswith('$') and not isinstance(value, dict)
            }
            self._apply_update(document, update, inserting=True)
            self.insert_one(document)
            return copy.deepcopy(document)
        return None

    def delete_one(self, query):
        deleted = self.find_one_and_delete(query)
        return SimpleNamespace(deleted_count=1 if deleted else 0)

    def delete_many(self, query):
        before = len(self.documents)
        self.documents = [document for document in self.documents if not self._matches(document, query)]
        return SimpleNamespace(deleted_count=before - len(self.documents))

    @staticmethod
    def _apply_update(document, update, inserting):
        for key, value in update.get('$set', {}).items():
            document[key] = copy.deepcopy(value)
        if inserting:
            for key, value in update.get('$setOnInsert', {}).items():
                document[key] = copy.deepcopy(value)
        for key in update.get('$unset', {}):
            document.pop(key, None)
        for key, value in update.get('$inc', {}).items():
            document[key] = document.get(key, 0) + value


class FakeIndexCollection:
    def __init__(self, indexes=None):
        self.indexes = copy.deepcopy(indexes or [{"name": "_id_", "key": {"_id": 1}, "unique": True}])

    def list_indexes(self):
        return copy.deepcopy(self.indexes)

    def drop_index(self, name):
        self.indexes = [index for index in self.indexes if index['name'] != name]

    def create_index(self, keys, name, **options):
        normalized = [(keys, 1)] if isinstance(keys, str) else list(keys)
        self.indexes.append({'name': name, 'key': dict(normalized), **options})
        return name


class FakeDatabase:
    def __init__(self):
        self.collections = {name: FakeIndexCollection() for name in db_indexes.INDEX_DEFINITIONS}

    def __getitem__(self, name):
        return self.collections[name]


class BackendSecurityTests(unittest.TestCase):
    def setUp(self):
        self.saved_globals = {
            name: getattr(backend, name)
            for name in (
                'users_collection', 'devices_collection', 'device_status_collection',
                'pages_collection', 'pairing_codes_collection', 'ALLOW_REGISTRATION',
                'DEVICE_AUTH_REQUIRED', 'ADMIN_BOOTSTRAP_TOKEN', 'DATA_DIR',
            )
        }
        backend.app.config.update(TESTING=True)
        backend.users_collection = FakeCollection()
        backend.devices_collection = FakeCollection()
        backend.device_status_collection = FakeCollection()
        backend.pages_collection = FakeCollection()
        backend.pairing_codes_collection = FakeCollection()
        backend.ALLOW_REGISTRATION = False
        backend.DEVICE_AUTH_REQUIRED = True
        backend.ADMIN_BOOTSTRAP_TOKEN = 'bootstrap-' + ('A' * 32)
        self.temp_directory = tempfile.TemporaryDirectory()
        backend.DATA_DIR = Path(self.temp_directory.name)
        self.client = backend.app.test_client()
        self.user_patch = patch.object(
            backend, 'get_current_user', return_value={'_id': 2, 'username': 'bob'}
        )
        self.user_patch.start()

    def tearDown(self):
        self.user_patch.stop()
        for name, value in self.saved_globals.items():
            setattr(backend, name, value)
        self.temp_directory.cleanup()

    def test_existing_owner_cannot_be_overwritten(self):
        backend.devices_collection = FakeCollection([
            {'deviceId': 'A1B2C3', 'owner': 'alice', 'claimed': True, 'imageVersion': 3}
        ])
        backend.pairing_codes_collection = FakeCollection([
            {'deviceId': 'A1B2C3', 'code': '123456', 'expiresAt': datetime.now(timezone.utc).replace(tzinfo=None) + timedelta(hours=1)}
        ])

        response = self.client.post('/api/devices/add', json={
            'deviceId': 'A1B2C3', 'pairingCode': '123456'
        })

        self.assertEqual(response.status_code, 403)
        self.assertEqual(backend.devices_collection.documents[0]['owner'], 'alice')
        self.assertEqual(len(backend.pairing_codes_collection.documents), 1)

    def test_pairing_code_is_required(self):
        backend.pairing_codes_collection = FakeCollection([
            {'deviceId': 'A1B2C3', 'code': '123456', 'expiresAt': datetime.now(timezone.utc).replace(tzinfo=None) + timedelta(hours=1)}
        ])

        response = self.client.post('/api/device/claim', json={'deviceId': 'A1B2C3'})

        self.assertEqual(response.status_code, 400)
        self.assertIn('pairing code', response.get_json()['error'].lower())

    def test_valid_pairing_code_is_consumed_once(self):
        backend.pairing_codes_collection = FakeCollection([
            {'deviceId': 'A1B2C3', 'code': '123456', 'expiresAt': datetime.now(timezone.utc).replace(tzinfo=None) + timedelta(hours=1)}
        ])

        response = self.client.post('/api/device/claim', json={
            'deviceId': 'A1B2C3', 'pairingCode': '123456'
        })

        self.assertEqual(response.status_code, 200)
        claimed = backend.devices_collection.find_one({'deviceId': 'A1B2C3'})
        self.assertEqual(claimed['owner'], 'bob')
        self.assertTrue(claimed['claimed'])
        self.assertEqual(backend.pairing_codes_collection.documents, [])

    def test_pairing_code_locks_after_repeated_failures(self):
        backend.pairing_codes_collection = FakeCollection([
            {
                'deviceId': 'A1B2C3', 'code': '123456',
                'expiresAt': datetime.now(timezone.utc).replace(tzinfo=None) + timedelta(hours=1),
                'failedAttempts': 0,
            }
        ])

        for _ in range(backend.PAIRING_MAX_FAILED_ATTEMPTS):
            response = self.client.post('/api/device/claim', json={
                'deviceId': 'A1B2C3', 'pairingCode': '000000'
            })
            self.assertEqual(response.status_code, 400)

        pairing = backend.pairing_codes_collection.find_one({'deviceId': 'A1B2C3'})
        self.assertGreater(pairing['lockedUntil'], datetime.now(timezone.utc).replace(tzinfo=None))
        correct_while_locked = self.client.post('/api/device/claim', json={
            'deviceId': 'A1B2C3', 'pairingCode': '123456'
        })
        self.assertEqual(correct_while_locked.status_code, 400)

    def test_additional_registration_is_closed_by_default(self):
        backend.users_collection = FakeCollection([
            {'_id': 1, 'username': 'admin', 'passwordHash': backend.hash_password('existing-pass')}
        ])

        response = self.client.post('/api/auth/register', json={
            'username': 'second', 'password': 'strong-pass'
        })

        self.assertEqual(response.status_code, 403)
        self.assertEqual(len(backend.users_collection.documents), 1)

    def test_first_registration_requires_bootstrap_token(self):
        response = self.client.post('/api/auth/register', json={
            'username': 'admin', 'password': 'strong-pass'
        })

        self.assertEqual(response.status_code, 403)
        self.assertEqual(backend.users_collection.documents, [])

    def test_first_registration_accepts_configured_bootstrap_token(self):
        response = self.client.post(
            '/api/auth/register',
            json={'username': 'admin', 'password': 'strong-pass'},
            headers={'X-Admin-Bootstrap-Token': backend.ADMIN_BOOTSTRAP_TOKEN},
        )

        self.assertEqual(response.status_code, 200)
        created = backend.users_collection.find_one({'username': 'admin'})
        self.assertTrue(created['passwordHash'].startswith('scrypt:'))
        self.assertNotIn('bootstrapToken', created)

    def test_login_stores_only_hashed_token_and_logout_revokes_it(self):
        password = 'correct-horse-battery-staple'
        backend.users_collection = FakeCollection([{
            '_id': 1,
            'username': 'alice',
            'passwordHash': backend.hash_password(password),
        }])

        logged_in = self.client.post('/api/auth/login', json={
            'username': 'alice', 'password': password,
        })
        self.assertEqual(logged_in.status_code, 200)
        token = logged_in.get_json()['token']
        stored = backend.users_collection.find_one({'username': 'alice'})
        self.assertNotIn('token', stored)
        self.assertEqual(stored['tokenHash'], backend.hash_token(token))
        self.assertNotEqual(stored['tokenHash'], token)

        # Exercise the real Bearer lookup instead of the owner stub used by
        # device-management tests in this class.
        self.user_patch.stop()
        try:
            authorized = self.client.get(
                '/api/auth/me', headers={'Authorization': f'Bearer {token}'},
            )
            logged_out = self.client.post(
                '/api/auth/logout', headers={'Authorization': f'Bearer {token}'},
            )
            revoked = self.client.get(
                '/api/auth/me', headers={'Authorization': f'Bearer {token}'},
            )
        finally:
            self.user_patch.start()

        self.assertEqual(authorized.status_code, 200)
        self.assertEqual(logged_out.status_code, 200)
        self.assertEqual(revoked.status_code, 401)

    def test_successful_login_migrates_legacy_password_hash(self):
        password = 'legacy-password'
        backend.users_collection = FakeCollection([{
            '_id': 1,
            'username': 'legacy',
            'passwordHash': hashlib.sha256(password.encode('utf-8')).hexdigest(),
        }])

        response = self.client.post('/api/auth/login', json={
            'username': 'legacy', 'password': password,
        })

        self.assertEqual(response.status_code, 200)
        migrated = backend.users_collection.find_one({'username': 'legacy'})
        self.assertTrue(migrated['passwordHash'].startswith('scrypt:'))

    def test_six_color_dimensions_are_fixed(self):
        response = self.client.post('/api/epd/process-sixcolor', json={
            'imageData': 'AAAA', 'width': 801, 'height': 480
        })

        self.assertEqual(response.status_code, 400)
        self.assertIn('800x480', response.get_json()['error'])

    def test_device_key_tofu_cannot_be_replaced(self):
        first_key = 'A' * 64
        second_key = 'B' * 64

        first = self.client.post(
            '/api/device/status',
            json={'deviceId': 'A1B2C3'},
            headers={'X-Device-Key': first_key},
        )
        rejected = self.client.post(
            '/api/device/status',
            json={'deviceId': 'A1B2C3'},
            headers={'X-Device-Key': second_key},
        )
        raw_rejected = self.client.get(
            '/api/epd/raw/A1B2C3',
            headers={'X-Device-Key': second_key},
        )

        self.assertEqual(first.status_code, 200)
        self.assertTrue(first.get_json()['success'])
        self.assertRegex(first.get_json()['pairingCode'], r'^\d{6}$')
        self.assertEqual(rejected.status_code, 401)
        self.assertEqual(raw_rejected.status_code, 401)
        stored = backend.device_status_collection.find_one({'deviceId': 'A1B2C3'})
        self.assertEqual(
            stored['deviceKeyHash'],
            hashlib.sha256(first_key.lower().encode('ascii')).hexdigest(),
        )
        self.assertGreater(
            stored['unclaimedExpiresAt'], datetime.now(timezone.utc).replace(tzinfo=None),
        )

    def test_device_write_locks_use_a_bounded_stable_pool(self):
        self.assertEqual(len(backend._device_write_locks), backend.DEVICE_WRITE_LOCK_STRIPES)
        self.assertIs(
            backend.get_device_write_lock('A1B2C3'),
            backend.get_device_write_lock('A1-B2-C3'),
        )

    def test_claimed_device_identity_is_exempt_from_unclaimed_ttl(self):
        device_key = 'A' * 64
        backend.devices_collection = FakeCollection([{
            'deviceId': 'A1B2C3', 'owner': 'bob', 'claimed': True,
        }])
        backend.device_status_collection = FakeCollection([{
            'deviceId': 'A1B2C3',
            'deviceKeyHash': hashlib.sha256(device_key.lower().encode('ascii')).hexdigest(),
            'unclaimedExpiresAt': datetime.now(timezone.utc).replace(tzinfo=None),
        }])

        response = self.client.post(
            '/api/device/status',
            json={'deviceId': 'A1B2C3'},
            headers={'X-Device-Key': device_key},
        )

        self.assertEqual(response.status_code, 200)
        stored = backend.device_status_collection.find_one({'deviceId': 'A1B2C3'})
        self.assertTrue(stored['everClaimed'])
        self.assertNotIn('unclaimedExpiresAt', stored)

    def test_compatibility_mode_only_bypasses_devices_without_a_registered_key(self):
        backend.DEVICE_AUTH_REQUIRED = False
        registered_key = 'C' * 64
        registered_hash = hashlib.sha256(registered_key.lower().encode('ascii')).hexdigest()
        backend.device_status_collection = FakeCollection([
            {'deviceId': 'C1D2E3', 'deviceKeyHash': registered_hash},
        ])

        legacy = self.client.post('/api/device/status', json={'deviceId': 'A1B2C3'})
        tofu_key = 'D' * 64
        tofu = self.client.post(
            '/api/device/status',
            json={'deviceId': 'B1C2D3'},
            headers={'X-Device-Key': tofu_key},
        )
        protected_without_key = self.client.post(
            '/api/device/status', json={'deviceId': 'C1D2E3'}
        )
        protected_with_wrong_key = self.client.post(
            '/api/device/status',
            json={'deviceId': 'C1D2E3'},
            headers={'X-Device-Key': 'E' * 64},
        )

        self.assertEqual(legacy.status_code, 200)
        self.assertEqual(tofu.status_code, 200)
        self.assertEqual(protected_without_key.status_code, 401)
        self.assertEqual(protected_with_wrong_key.status_code, 401)
        stored = backend.device_status_collection.find_one({'deviceId': 'B1C2D3'})
        self.assertEqual(
            stored['deviceKeyHash'],
            hashlib.sha256(tofu_key.lower().encode('ascii')).hexdigest(),
        )

    def test_device_telemetry_strips_control_characters(self):
        response = self.client.post(
            '/api/device/status',
            json={
                'deviceId': 'A1B2C3',
                'ip': '192.168.1.2\nFORGED',
                'wakeType': 'manual',
                'wakeCause': 'GPIO\r\nFORGED',
            },
            headers={'X-Device-Key': 'A' * 64},
        )

        self.assertEqual(response.status_code, 200)
        stored = backend.device_status_collection.find_one({'deviceId': 'A1B2C3'})
        self.assertNotRegex(stored['ip'], r'[\r\n]')
        self.assertNotRegex(stored['lastWakeCause'], r'[\r\n]')

    def test_authenticated_update_result_is_stored_and_exposed(self):
        device_key = 'A' * 64
        backend.devices_collection = FakeCollection([{
            'deviceId': 'A1B2C3', 'owner': 'bob', 'claimed': True, 'imageVersion': 8,
        }])
        backend.device_status_collection = FakeCollection([{
            'deviceId': 'A1B2C3',
            'deviceKeyHash': hashlib.sha256(device_key.lower().encode('ascii')).hexdigest(),
        }])
        payload = {
            'deviceId': 'A1B2C3',
            'firmwareVersion': '3.1.0',
            'firmwareBuild': 'Aug 03 2026 12:00:00',
            'resetReason': 'DEEPSLEEP',
            'localImageVersion': 8,
            'targetImageVersion': 8,
            'updateAttemptId': '0123456789abcdef',
            'updateResult': 'success',
            'updateStage': 'done',
            'updateError': 'none',
            'updateDurationMs': 123456,
            'gpio0StuckLow': False,
        }

        result = self.client.post(
            '/api/device/update-result', json=payload,
            headers={'X-Device-Key': device_key},
        )
        devices = self.client.get('/api/devices')

        self.assertEqual(result.status_code, 200)
        self.assertTrue(result.get_json()['diagnosticAccepted'])
        stored = backend.device_status_collection.find_one({'deviceId': 'A1B2C3'})
        self.assertEqual(stored['lastUpdateResult'], 'success')
        self.assertEqual(stored['lastUpdateStage'], 'done')
        self.assertEqual(stored['localImageVersion'], 8)
        self.assertEqual(devices.status_code, 200)
        exposed = devices.get_json()['devices'][0]
        self.assertEqual(exposed['firmwareVersion'], '3.1.0')
        self.assertEqual(exposed['lastUpdateDurationMs'], 123456)

    def test_presence_marks_device_offline_after_first_missed_automatic_wake(self):
        now_ms = 1_800_000_000_000
        sleep_interval_seconds = 60 * 60
        backend.devices_collection = FakeCollection([{
            'deviceId': 'A1B2C3',
            'owner': 'bob',
            'claimed': True,
            'imageVersion': 8,
            'activeContentMode': 'template',
            'activeTemplateId': 'nameplate',
            'templateConfig': {'sleepIntervalSeconds': sleep_interval_seconds},
        }])

        cases = [
            ('active wake window', 60, True, False, 1),
            ('normal sleep', 10 * 60, False, True, 1),
            ('first wake grace', 64 * 60, False, True, 1),
            ('first wake missed', 66 * 60, False, False, 1),
        ]

        with patch.object(backend.time, 'time', return_value=now_ms / 1000):
            for name, age_seconds, expected_online, expected_sleeping, expected_wake_number in cases:
                with self.subTest(name=name):
                    last_seen = now_ms - age_seconds * 1000
                    backend.device_status_collection = FakeCollection([{
                        'deviceId': 'A1B2C3',
                        'lastSeen': last_seen,
                        'currentSleepSeconds': sleep_interval_seconds,
                    }])

                    response = self.client.get('/api/devices')
                    self.assertEqual(response.status_code, 200)
                    exposed = response.get_json()['devices'][0]
                    self.assertEqual(exposed['online'], expected_online)
                    self.assertEqual(exposed['sleeping'], expected_sleeping)
                    self.assertEqual(
                        exposed['estimatedNextAutoWakeAt'],
                        last_seen + expected_wake_number * sleep_interval_seconds * 1000,
                    )

    def test_update_result_rejects_unknown_stage_without_overwriting_status(self):
        device_key = 'A' * 64
        backend.device_status_collection = FakeCollection([{
            'deviceId': 'A1B2C3',
            'deviceKeyHash': hashlib.sha256(device_key.lower().encode('ascii')).hexdigest(),
            'lastUpdateResult': 'success',
        }])
        payload = {
            'deviceId': 'A1B2C3',
            'firmwareVersion': '3.1.0', 'firmwareBuild': 'Aug 03 2026 12:00:00',
            'resetReason': 'DEEPSLEEP', 'localImageVersion': 8,
            'targetImageVersion': 9, 'updateAttemptId': '0123456789abcdef',
            'updateResult': 'failed', 'updateStage': 'arbitrary_stage',
            'updateError': 'download_http', 'updateDurationMs': 100,
            'gpio0StuckLow': False,
        }
        response = self.client.post(
            '/api/device/update-result', json=payload,
            headers={'X-Device-Key': device_key},
        )

        self.assertEqual(response.status_code, 400)
        stored = backend.device_status_collection.find_one({'deviceId': 'A1B2C3'})
        self.assertEqual(stored['lastUpdateResult'], 'success')

    def test_status_accepts_interrupted_update_diagnostic_for_next_wake_replay(self):
        device_key = 'A' * 64
        response = self.client.post(
            '/api/device/status',
            json={
                'deviceId': 'A1B2C3', 'diagnosticPresent': True,
                'firmwareVersion': '3.1.0', 'resetReason': 'BROWNOUT',
                'localImageVersion': 7, 'targetImageVersion': 8,
                'updateAttemptId': 'fedcba9876543210',
                'updateResult': 'interrupted', 'updateStage': 'epd_refresh',
                'updateError': 'interrupted', 'updateDurationMs': 8000,
                'gpio0StuckLow': False,
            },
            headers={'X-Device-Key': device_key},
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.get_json()['diagnosticAccepted'])
        stored = backend.device_status_collection.find_one({'deviceId': 'A1B2C3'})
        self.assertEqual(stored['lastUpdateResult'], 'interrupted')
        self.assertEqual(stored['lastUpdateStage'], 'epd_refresh')
        self.assertEqual(stored['resetReason'], 'BROWNOUT')

    def test_owner_can_open_one_short_device_key_reset_window(self):
        old_key = 'A' * 64
        new_key = 'B' * 64
        backend.devices_collection = FakeCollection([
            {'deviceId': 'A1B2C3', 'owner': 'bob', 'claimed': True, 'imageVersion': 0},
        ])
        backend.device_status_collection = FakeCollection([{
            'deviceId': 'A1B2C3',
            'deviceKeyHash': hashlib.sha256(old_key.lower().encode('ascii')).hexdigest(),
        }])

        opened = self.client.post('/api/device/auth/reset', json={'deviceId': 'A1B2C3'})
        rotated = self.client.post(
            '/api/device/status',
            json={'deviceId': 'A1B2C3'},
            headers={'X-Device-Key': new_key},
        )
        old_rejected = self.client.post(
            '/api/device/status',
            json={'deviceId': 'A1B2C3'},
            headers={'X-Device-Key': old_key},
        )

        self.assertEqual(opened.status_code, 200)
        self.assertEqual(rotated.status_code, 200)
        self.assertEqual(old_rejected.status_code, 401)
        stored = backend.device_status_collection.find_one({'deviceId': 'A1B2C3'})
        self.assertEqual(
            stored['deviceKeyHash'],
            hashlib.sha256(new_key.lower().encode('ascii')).hexdigest(),
        )
        self.assertNotIn('deviceKeyResetUntil', stored)

    def test_invalid_stored_image_is_not_advertised_or_served(self):
        device_key = 'A' * 64
        device_hash = hashlib.sha256(device_key.lower().encode('ascii')).hexdigest()
        backend.devices_collection = FakeCollection([{
            'deviceId': 'A1B2C3',
            'owner': 'bob',
            'claimed': True,
            'imageVersion': 2,
            'imageSizeChars': backend.EPD_EXPECTED_CHARS,
            'imageSha256': hashlib.sha256(b'invalid').hexdigest(),
            'activeContentMode': 'image',
        }])
        backend.device_status_collection = FakeCollection([{
            'deviceId': 'A1B2C3', 'deviceKeyHash': device_hash,
        }])
        image_path = backend.get_device_image_path('A1B2C3', create_parent=True)
        image_path.write_text('invalid', encoding='utf-8')

        status = self.client.post(
            '/api/device/status',
            json={'deviceId': 'A1B2C3'},
            headers={'X-Device-Key': device_key},
        )
        download = self.client.get(
            '/api/epd/raw/A1B2C3', headers={'X-Device-Key': device_key}
        )

        self.assertEqual(status.status_code, 200)
        self.assertNotIn('imageUrl', status.get_json())
        self.assertEqual(download.status_code, 503)

    def test_exact_length_stored_image_with_invalid_symbols_is_rejected(self):
        device_key = 'A' * 64
        device_hash = hashlib.sha256(device_key.lower().encode('ascii')).hexdigest()
        backend.devices_collection = FakeCollection([{
            'deviceId': 'A1B2C3',
            'owner': 'bob',
            'claimed': True,
            'imageVersion': 2,
            'imageSizeChars': backend.EPD_EXPECTED_CHARS,
            'activeContentMode': 'image',
        }])
        backend.device_status_collection = FakeCollection([{
            'deviceId': 'A1B2C3', 'deviceKeyHash': device_hash,
        }])
        image_path = backend.get_device_image_path('A1B2C3', create_parent=True)
        image_path.write_bytes(b'q' * backend.EPD_EXPECTED_CHARS)

        status = self.client.post(
            '/api/device/status',
            json={'deviceId': 'A1B2C3'},
            headers={'X-Device-Key': device_key},
        )
        download = self.client.get(
            '/api/epd/raw/A1B2C3', headers={'X-Device-Key': device_key}
        )

        self.assertEqual(status.status_code, 200)
        self.assertNotIn('imageUrl', status.get_json())
        self.assertEqual(download.status_code, 503)

    def test_public_image_url_uses_configured_origin(self):
        with patch.object(backend.Config, 'PUBLIC_BASE_URL', 'https://epd.example.test'):
            self.assertEqual(
                backend.build_raw_image_url('A1B2C3', 7),
                'https://epd.example.test/api/epd/raw/A1B2C3?v=7',
            )

    def test_versioned_raw_download_serves_exact_retained_version(self):
        device_key = 'A' * 64
        payloads = {}
        for version in range(1, 7):
            payload = chr(ord('a') + version) * backend.EPD_EXPECTED_CHARS
            payloads[version] = payload
            self.assertTrue(backend.save_device_image('A1B2C3', payload, version))
        backend.devices_collection = FakeCollection([{
            'deviceId': 'A1B2C3', 'owner': 'bob', 'claimed': True,
            'imageVersion': 6,
            'imageSizeChars': backend.EPD_EXPECTED_CHARS,
            'imageSha256': hashlib.sha256(payloads[6].encode('ascii')).hexdigest(),
        }])
        backend.device_status_collection = FakeCollection([{
            'deviceId': 'A1B2C3',
            'deviceKeyHash': hashlib.sha256(device_key.lower().encode('ascii')).hexdigest(),
        }])

        retained = self.client.get(
            '/api/epd/raw/A1B2C3?v=2', headers={'X-Device-Key': device_key},
        )
        expired = self.client.get(
            '/api/epd/raw/A1B2C3?v=1', headers={'X-Device-Key': device_key},
        )
        current = self.client.get(
            '/api/epd/raw/A1B2C3?v=6', headers={'X-Device-Key': device_key},
        )

        self.assertEqual(retained.status_code, 200)
        self.assertEqual(retained.data, payloads[2].encode('ascii'))
        self.assertEqual(retained.headers['X-EPD-Version'], '2')
        self.assertEqual(expired.status_code, 409)
        self.assertEqual(expired.get_json()['errorCode'], 'IMAGE_VERSION_EXPIRED')
        self.assertEqual(current.status_code, 200)
        self.assertEqual(current.data, payloads[6].encode('ascii'))
        retained.close()
        current.close()

    def test_first_versioned_publish_preserves_legacy_latest_snapshot(self):
        legacy_payload = 'h' * backend.EPD_EXPECTED_CHARS
        new_payload = 'i' * backend.EPD_EXPECTED_CHARS
        legacy_hash = hashlib.sha256(legacy_payload.encode('ascii')).hexdigest()
        latest_path = backend.get_device_image_path('A1B2C3', create_parent=True)
        latest_path.write_text(legacy_payload, encoding='utf-8')

        saved = backend.save_device_image(
            'A1B2C3',
            new_payload,
            8,
            previous_version=7,
            previous_sha256=legacy_hash,
        )

        self.assertTrue(saved)
        self.assertEqual(
            backend.get_device_version_image_path('A1B2C3', 7).read_text(encoding='utf-8'),
            legacy_payload,
        )
        self.assertEqual(
            backend.get_device_version_image_path('A1B2C3', 8).read_text(encoding='utf-8'),
            new_payload,
        )
        self.assertEqual(latest_path.read_text(encoding='utf-8'), new_payload)

    def test_weather_and_quote_require_login(self):
        with patch.object(backend, 'get_current_user', return_value=None):
            weather = self.client.get('/api/weather?city=Shanghai')
            quote = self.client.get('/api/quote')
        self.assertEqual(weather.status_code, 401)
        self.assertEqual(quote.status_code, 401)

    def test_legacy_page_without_device_scope_is_fail_closed(self):
        backend.pages_collection = FakeCollection([{
            'pageId': 'legacy-page', 'name': 'orphan', 'data': {},
        }])

        fetched = self.client.get('/api/pages/legacy-page')
        deleted = self.client.delete('/api/pages/legacy-page')

        self.assertEqual(fetched.status_code, 403)
        self.assertEqual(deleted.status_code, 403)
        self.assertIsNotNone(backend.pages_collection.find_one({'pageId': 'legacy-page'}))

    def test_page_list_only_returns_current_owner_and_ownerless_pages(self):
        backend.devices_collection = FakeCollection([{
            'deviceId': 'A1B2C3', 'owner': 'bob', 'claimed': True,
        }])
        now = datetime.now(timezone.utc).replace(tzinfo=None)
        backend.pages_collection = FakeCollection([
            {
                'pageId': 'bob-page', 'deviceId': 'A1B2C3', 'owner': 'bob',
                'name': 'Bob page', 'updatedAt': now,
            },
            {
                'pageId': 'alice-page', 'deviceId': 'A1B2C3', 'owner': 'alice',
                'name': 'Alice page', 'updatedAt': now - timedelta(seconds=1),
            },
            {
                'pageId': 'legacy-page', 'deviceId': 'A1B2C3',
                'name': 'Legacy page', 'updatedAt': now - timedelta(seconds=2),
            },
            {
                'pageId': 'null-owner-page', 'deviceId': 'A1B2C3', 'owner': None,
                'name': 'Null owner page', 'updatedAt': now - timedelta(seconds=3),
            },
        ])

        response = self.client.get('/api/pages/list/A1B2C3')

        self.assertEqual(response.status_code, 200)
        page_ids = {page['pageId'] for page in response.get_json()['pages']}
        self.assertEqual(page_ids, {'bob-page', 'legacy-page', 'null-owner-page'})

    def test_page_lookup_prefers_the_current_owner_when_ids_collide(self):
        backend.devices_collection = FakeCollection([{
            'deviceId': 'A1B2C3', 'owner': 'bob', 'claimed': True,
        }])
        backend.pages_collection = FakeCollection([
            {
                '_id': 1, 'pageId': 'shared-page', 'deviceId': 'D1E2F3',
                'owner': 'alice', 'name': 'Alice page', 'data': {},
            },
            {
                '_id': 2, 'pageId': 'shared-page', 'deviceId': 'A1B2C3',
                'owner': 'bob', 'name': 'Bob page', 'data': {},
            },
        ])

        fetched = self.client.get('/api/pages/shared-page')
        deleted = self.client.delete('/api/pages/shared-page')

        self.assertEqual(fetched.status_code, 200)
        self.assertEqual(fetched.get_json()['page']['name'], 'Bob page')
        self.assertEqual(deleted.status_code, 200)
        self.assertIsNotNone(backend.pages_collection.find_one({'_id': 1}))
        self.assertIsNone(backend.pages_collection.find_one({'_id': 2}))

    def test_new_page_ids_use_full_uuid_entropy(self):
        backend.devices_collection = FakeCollection([{
            'deviceId': 'A1B2C3', 'owner': 'bob', 'claimed': True,
        }])

        response = self.client.post('/api/pages/save', json={
            'deviceId': 'A1B2C3', 'name': 'New page', 'data': {},
        })

        self.assertEqual(response.status_code, 200)
        self.assertRegex(response.get_json()['pageId'], r'^[0-9a-f]{32}$')

    def test_page_thumbnail_size_is_limited(self):
        response = self.client.post('/api/pages/save', json={
            'deviceId': 'A1B2C3',
            'name': 'test',
            'type': 'custom',
            'data': {},
            'thumbnail': 'x' * (backend.PAGE_MAX_THUMBNAIL_BYTES + 1),
        })
        self.assertEqual(response.status_code, 413)

    def test_legacy_page_device_id_is_normalized_on_save(self):
        backend.devices_collection = FakeCollection([{
            'deviceId': 'A1B2C3', 'owner': 'bob', 'claimed': True,
        }])
        backend.pages_collection = FakeCollection([{
            'pageId': 'legacy-page', 'deviceId': 'a1-b2-c3',
            'name': 'old', 'type': 'custom', 'data': {},
        }])

        response = self.client.post('/api/pages/save', json={
            'deviceId': 'A1B2C3',
            'pageId': 'legacy-page',
            'name': 'new',
            'type': 'custom',
            'data': {},
        })

        self.assertEqual(response.status_code, 200)
        stored = backend.pages_collection.find_one({'pageId': 'legacy-page'})
        self.assertEqual(stored['deviceId'], 'A1B2C3')
        self.assertEqual(stored['owner'], 'bob')

    def test_nameplate_parsers_preserve_intentional_duplicate_names(self):
        expected = ['张伟', '张伟']
        self.assertEqual(backend.parse_nameplate_names({'names': expected}), expected)
        self.assertEqual(backend.parse_nameplate_names({'text': '张伟\n张伟'}), expected)
        self.assertEqual(backend.parse_nameplate_names_from_text('张伟\n张伟'), expected)
        self.assertEqual(
            backend.build_nameplate_parse_result(expected, {})['names'], expected
        )

    def test_nameplate_parser_preserves_english_full_names(self):
        expected = ['Alexander Montgomery', 'Mary-Jane Watson']
        self.assertEqual(backend.parse_nameplate_names({'names': expected}), expected)
        self.assertEqual(
            backend.parse_nameplate_names_from_text('\n'.join(expected)), expected
        )
        self.assertEqual(backend._clean_nameplate_candidate('张三 技术专家'), '张三')

    def test_nameplate_parser_matches_title_and_company_per_person(self):
        people = backend.parse_nameplate_people_from_text(
            '姓名\t公司名称\t职位\n'
            '张伟\t甲公司\t总经理\n'
            '张伟\t乙公司\t副总经理\n'
            'Alexander Montgomery\tPheno\tTechnical Expert'
        )

        self.assertEqual(people, [
            {'name': '张伟', 'title': '总经理', 'subtitle': '甲公司'},
            {'name': '张伟', 'title': '副总经理', 'subtitle': '乙公司'},
            {
                'name': 'Alexander Montgomery',
                'title': 'Technical Expert',
                'subtitle': 'Pheno',
            },
        ])
        parsed = backend.build_nameplate_parse_result(people, {})
        self.assertEqual(parsed['names'], ['张伟', '张伟', 'Alexander Montgomery'])
        self.assertEqual(parsed['people'], people)

    def test_nameplate_person_joins_multiple_titles_and_companies_with_slashes(self):
        person = backend.normalize_nameplate_person({
            'name': '刘清侠',
            'title': ['院长', '院士'],
            'position': '院长',
            'subtitle': '深圳技术大学聚龙学院、加拿大工程院',
        })

        self.assertEqual(person, {
            'name': '刘清侠',
            'title': '院长 / 院士',
            'subtitle': '深圳技术大学聚龙学院 / 加拿大工程院',
        })

        second_person = backend.normalize_nameplate_person({
            'name': '张国芳',
            '职位': '主席；董事长',
            '公司': '坪山区工商业联合会 / 六和集团',
        })
        self.assertEqual(second_person['title'], '主席 / 董事长')
        self.assertEqual(second_person['subtitle'], '坪山区工商业联合会 / 六和集团')

    def test_nameplate_preview_renders_without_publishing(self):
        backend.devices_collection = FakeCollection([{
            'deviceId': 'A1B2C3', 'owner': 'bob', 'claimed': True, 'imageVersion': 8,
        }])
        preview_before = copy.deepcopy(backend.devices_collection.documents)

        with patch.object(
            backend, 'render_template_with_preview',
            return_value={'previewImage': 'preview-base64'},
        ) as render_mock:
            response = self.client.post('/api/nameplates/preview', json={
                'name': 'Alexander Montgomery',
                'templateConfig': {
                    'backgroundStyle': 'formal_red',
                    'title': 'Technical Expert',
                    'subtitle': 'Pheno',
                },
            })

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.get_json()['previewImage'], 'preview-base64')
        rendered_config = render_mock.call_args.args[1]
        self.assertEqual(rendered_config['name'], 'Alexander Montgomery')
        self.assertEqual(rendered_config['title'], 'Technical Expert')
        self.assertEqual(backend.devices_collection.documents, preview_before)

    def test_nameplate_preview_keeps_company_center_when_person_company_changes(self):
        with patch.object(
            backend, 'render_template_with_preview',
            return_value={'previewImage': 'preview-base64'},
        ) as render_mock:
            response = self.client.post('/api/nameplates/preview', json={
                'person': {
                    'name': '罗宇航',
                    'title': '副局长',
                    'subtitle': '中山市投资促进局',
                },
                'templateConfig': {
                    'backgroundStyle': 'formal_red',
                    'subtitle': 'Advanced Quantum Microelectronics & Semiconductor',
                    'companyX': 260,
                    'companyPositionMode': 'custom',
                },
            })

        self.assertEqual(response.status_code, 200)
        rendered_config = render_mock.call_args.args[1]
        self.assertEqual(rendered_config['subtitle'], '中山市投资促进局')
        self.assertEqual(
            rendered_config['companyReferenceText'],
            'Advanced Quantum Microelectronics & Semi',
        )

    def test_nameplate_parse_accepts_standard_comma_separated_text_without_ai(self):
        with patch.object(backend, 'get_nameplate_ai_api_key', return_value=''):
            response = self.client.post('/api/nameplates/parse', data={
                'text': '张三，李四，Alexander Montgomery',
                'templateConfig': '{}',
            })

        self.assertEqual(response.status_code, 200)
        parsed = response.get_json()['parsed']
        self.assertFalse(parsed['aiUsed'])
        self.assertEqual(
            parsed['names'], ['张三', '李四', 'Alexander Montgomery']
        )

    def test_nameplate_ai_output_accepts_minimax_fenced_object_variant(self):
        output = """```json
{
  "template": {"backgroundStyle": "formal_red", "title": "嘉宾"},
  "names": [{"name": "张三"}, {"name": "李四"}],
  "warnings": ["请核对\\'张三\\'"]
}
```"""

        parsed = backend._parse_nameplate_ai_output(output)

        self.assertEqual(parsed['names'], ['张三', '李四'])
        self.assertEqual(parsed['templateConfig']['title'], '嘉宾')
        self.assertEqual(parsed['warnings'], ["请核对'张三'"])
        self.assertEqual(parsed['sourceSummary'], 'AI解析')

    def test_nameplate_ai_output_preserves_per_person_details(self):
        parsed = backend._parse_nameplate_ai_output(json.dumps({
            'people': [
                {'name': '张三', 'title': '总经理', 'subtitle': '甲公司'},
                {'name': '李四', 'title': '', 'subtitle': '乙公司'},
            ],
            'templateConfig': {'title': '嘉宾', 'subtitle': '默认公司'},
            'warnings': [],
            'sourceSummary': '表格解析',
        }, ensure_ascii=False))

        self.assertEqual(parsed['names'], ['张三', '李四'])
        self.assertEqual(parsed['people'][0]['title'], '总经理')
        self.assertEqual(parsed['people'][1]['subtitle'], '乙公司')

    def test_nameplate_ai_retry_shares_one_sub_120_second_budget(self):
        logo_data_url = make_test_logo_data_url()
        first_response = SimpleNamespace(
            status_code=400,
            text='response_format is unsupported',
        )
        second_response = SimpleNamespace(
            status_code=200,
            text='',
            json=lambda: {
                'choices': [{
                    'message': {
                        'content': json.dumps({
                            'names': ['张伟', '张伟'],
                            'templateConfig': {},
                            'warnings': [],
                            'sourceSummary': 'test',
                        }, ensure_ascii=False),
                    },
                }],
            },
        )

        with (
            patch.object(backend, 'get_nameplate_ai_api_key', return_value='test-key'),
            patch.dict(os.environ, {'NAMEPLATE_AI_API_MODE': 'chat_completions'}),
            patch.object(backend.time, 'monotonic', side_effect=[100.0, 100.0, 165.0]),
            patch.object(
                backend.requests, 'post', side_effect=[first_response, second_response]
            ) as post_mock,
        ):
            result = backend.call_openai_nameplate_parser('', [], {
                'backgroundStyle': 'formal_green',
                'title': '统一职位',
                'subtitle': '统一公司',
                'sleepIntervalSeconds': 21600,
                'logoDataUrl': logo_data_url,
                'logoFileName': 'event-logo.png',
                'logoX': 21,
                'logoY': 34,
                'companyX': 412,
                'companyPositionMode': 'custom',
            })

        timeouts = [call.kwargs['timeout'] for call in post_mock.call_args_list]
        self.assertLess(backend.NAMEPLATE_AI_TOTAL_BUDGET_SECONDS, 120)
        self.assertAlmostEqual(sum(timeouts[0]), 90.0)
        self.assertAlmostEqual(sum(timeouts[1]), 25.0)
        self.assertEqual(result['names'], ['张伟', '张伟'])
        self.assertEqual(result['templateConfig']['logoDataUrl'], logo_data_url)
        self.assertEqual(result['templateConfig']['logoX'], 21)
        self.assertEqual(result['templateConfig']['logoY'], 34)
        self.assertEqual(result['templateConfig']['companyX'], 412)
        self.assertEqual(result['templateConfig']['companyPositionMode'], 'custom')
        self.assertEqual(result['templateConfig']['backgroundStyle'], 'formal_green')
        self.assertEqual(result['templateConfig']['title'], '统一职位')
        self.assertEqual(result['templateConfig']['subtitle'], '统一公司')
        self.assertEqual(result['templateConfig']['sleepIntervalSeconds'], 21600)

    def test_nameplate_ai_prompt_requires_all_roles_and_companies_with_slashes(self):
        response = SimpleNamespace(
            status_code=200,
            text='',
            json=lambda: {
                'choices': [{
                    'message': {
                        'content': json.dumps({
                            'people': [{
                                'name': '刘清侠',
                                'title': '院长 / 院士',
                                'subtitle': '深圳技术大学聚龙学院 / 加拿大工程院',
                            }],
                            'templateConfig': {},
                            'warnings': [],
                            'sourceSummary': 'test',
                        }, ensure_ascii=False),
                    },
                }],
            },
        )

        with (
            patch.object(backend, 'get_nameplate_ai_api_key', return_value='test-key'),
            patch.dict(os.environ, {'NAMEPLATE_AI_API_MODE': 'chat_completions'}),
            patch.object(backend.requests, 'post', return_value=response) as post_mock,
        ):
            result = backend.call_openai_nameplate_parser('刘清侠名单信息', [], {})

        prompt = post_mock.call_args.kwargs['json']['messages'][0]['content'][0]['text']
        self.assertIn('不判断主次、不丢弃', prompt)
        self.assertIn('院长 / 院士', prompt)
        self.assertIn('多个职位或多个单位不属于不确定情况', prompt)
        self.assertEqual(result['people'][0]['title'], '院长 / 院士')
        self.assertEqual(
            result['people'][0]['subtitle'],
            '深圳技术大学聚龙学院 / 加拿大工程院',
        )
        self.assertEqual(result['warnings'], [])

    def test_nameplate_dispatch_deadline_reports_unprocessed_devices(self):
        backend.devices_collection = FakeCollection([
            {
                'deviceId': device_id,
                'deviceName': f'Device {index}',
                'owner': 'bob',
                'claimed': True,
                'imageVersion': 0,
            }
            for index, device_id in enumerate(('A1B2C3', 'D4E5F6', '112233'), start=1)
        ])
        epd_data = 'a' * backend.EPD_EXPECTED_CHARS

        with (
            patch.object(backend.time, 'monotonic', side_effect=[0.0, 1.0, 91.0]),
            patch.object(
                backend, 'render_template_with_preview',
                return_value={'epdData': epd_data},
            ),
            patch.object(backend, 'save_device_image', return_value=True),
        ):
            response = self.client.post('/api/nameplates/dispatch', json={
                'names': ['张一', '张二', '张三'],
                'deviceIds': ['A1B2C3', 'D4E5F6', '112233'],
            })

        payload = response.get_json()
        self.assertEqual(response.status_code, 200)
        self.assertTrue(payload['deadlineReached'])
        self.assertEqual(payload['processedCount'], 1)
        self.assertEqual(payload['assignedCount'], 1)
        self.assertEqual(
            [item['deviceId'] for item in payload['unprocessedDevices']],
            ['D4E5F6', '112233'],
        )
        self.assertTrue(all(item['unprocessed'] for item in payload['failed']))

    def test_nameplate_dispatch_count_is_limited_before_rendering(self):
        response = self.client.post('/api/nameplates/dispatch', json={
            'names': ['张三'] * (backend.NAMEPLATE_MAX_NAMES + 1),
        })
        self.assertEqual(response.status_code, 400)

    def test_nameplate_dispatch_prefers_per_person_title_and_company(self):
        backend.devices_collection = FakeCollection([
            {
                'deviceId': device_id,
                'deviceName': f'Device {index}',
                'owner': 'bob',
                'claimed': True,
                'imageVersion': 0,
            }
            for index, device_id in enumerate(('A1B2C3', 'D4E5F6'), start=1)
        ])
        epd_data = 'a' * backend.EPD_EXPECTED_CHARS

        with (
            patch.object(
                backend, 'render_template_with_preview',
                return_value={'epdData': epd_data},
            ) as render_mock,
            patch.object(backend, 'save_device_image', return_value=True),
        ):
            response = self.client.post('/api/nameplates/dispatch', json={
                'people': [
                    {'name': '张三', 'title': '总经理', 'subtitle': '甲公司'},
                    {'name': '李四', 'title': '', 'subtitle': ''},
                ],
                'deviceIds': ['A1B2C3', 'D4E5F6'],
                'templateConfig': {
                    'title': '统一嘉宾',
                    'subtitle': '统一公司',
                },
            })

        self.assertEqual(response.status_code, 200)
        rendered_configs = [call.args[1] for call in render_mock.call_args_list]
        self.assertEqual(rendered_configs[0]['title'], '总经理')
        self.assertEqual(rendered_configs[0]['subtitle'], '甲公司')
        self.assertEqual(rendered_configs[0]['companyReferenceText'], '统一公司')
        self.assertEqual(rendered_configs[1]['title'], '统一嘉宾')
        self.assertEqual(rendered_configs[1]['subtitle'], '统一公司')
        assignments = response.get_json()['assignments']
        self.assertEqual(assignments[0]['title'], '总经理')
        self.assertEqual(assignments[1]['subtitle'], '统一公司')

    def test_xlsx_expansion_is_limited_before_openpyxl_parses_it(self):
        with tempfile.SpooledTemporaryFile() as workbook:
            with zipfile.ZipFile(workbook, 'w', compression=zipfile.ZIP_DEFLATED) as archive:
                archive.writestr('xl/worksheets/sheet1.xml', b'A' * 2048)
            workbook.seek(0)
            workbook_bytes = workbook.read()

        with patch.object(backend, 'NAMEPLATE_MAX_XLSX_UNCOMPRESSED_BYTES', 1024):
            text, warnings = backend.extract_spreadsheet_text(workbook_bytes, 'names.xlsx')

        self.assertEqual(text, '')
        self.assertTrue(any('解压后体积' in warning for warning in warnings))

    def test_docx_extracts_body_paragraphs_and_tables_in_order(self):
        document = Document()
        document.add_paragraph('中山市交流会名单')
        first_table = document.add_table(rows=1, cols=3)
        first_table.rows[0].cells[0].text = '姓名'
        first_table.rows[0].cells[1].text = '来访嘉宾'
        first_table.rows[0].cells[2].text = '职位'
        for name, company, title in (
            ('张伟', '甲公司', '总经理'),
            ('张伟', '乙公司', '副总经理'),
        ):
            cells = first_table.add_row().cells
            cells[0].text = name
            cells[1].text = company
            cells[2].text = title
        document.add_paragraph('第二组')
        second_table = document.add_table(rows=1, cols=2)
        second_table.rows[0].cells[0].text = 'Alexander Montgomery'
        second_table.rows[0].cells[1].text = 'Technical Expert'

        text, images, warnings, image_bytes = backend.extract_docx_content(
            serialize_test_docx(document), '名单.docx', 8, 16 * 1024 * 1024,
        )

        self.assertLess(text.index('中山市交流会名单'), text.index('姓名\t来访嘉宾\t职位'))
        self.assertLess(text.index('张伟\t甲公司'), text.index('张伟\t乙公司'))
        self.assertLess(text.index('第二组'), text.index('Alexander Montgomery'))
        self.assertEqual(text.count('张伟\t'), 2)
        self.assertEqual(images, [])
        self.assertEqual(warnings, [])
        self.assertEqual(image_bytes, 0)

    def test_docx_extracts_embedded_images_and_enforces_image_limit(self):
        document = Document()
        document.add_paragraph('图片名单')
        image = Image.new('RGB', (40, 20), (255, 255, 255))
        image_buffer = io.BytesIO()
        image.save(image_buffer, format='PNG')
        image_buffer.seek(0)
        document.add_picture(image_buffer)
        raw_docx = serialize_test_docx(document)

        text, images, warnings, image_bytes = backend.extract_docx_content(
            raw_docx, '图片名单.docx', 8, 16 * 1024 * 1024,
        )
        self.assertIn('图片名单', text)
        self.assertEqual(len(images), 1)
        self.assertEqual(images[0]['mimeType'], 'image/png')
        self.assertTrue(images[0]['dataUrl'].startswith('data:image/png;base64,'))
        self.assertGreater(image_bytes, 0)
        self.assertEqual(warnings, [])

        _, limited_images, limited_warnings, limited_bytes = backend.extract_docx_content(
            raw_docx, '图片名单.docx', 0, 0,
        )
        self.assertEqual(limited_images, [])
        self.assertEqual(limited_bytes, 0)
        self.assertTrue(any('数量或总大小限制' in item for item in limited_warnings))

    def test_docx_parse_without_ai_preserves_text_names_and_does_not_publish(self):
        document = Document()
        for name in ('张伟', '张伟', 'Alexander Montgomery'):
            document.add_paragraph(name)
        image = Image.new('RGB', (20, 20), (255, 255, 255))
        image_buffer = io.BytesIO()
        image.save(image_buffer, format='PNG')
        image_buffer.seek(0)
        document.add_picture(image_buffer)
        raw_docx = serialize_test_docx(document)
        devices_before = copy.deepcopy(backend.devices_collection.documents)

        with patch.object(backend, 'get_nameplate_ai_api_key', return_value=''):
            response = self.client.post('/api/nameplates/parse', data={
                'files': (
                    io.BytesIO(raw_docx),
                    'names.docx',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                ),
                'templateConfig': '{}',
            }, content_type='multipart/form-data')

        self.assertEqual(response.status_code, 200)
        parsed = response.get_json()['parsed']
        self.assertFalse(parsed['aiUsed'])
        self.assertEqual(parsed['names'], ['张伟', '张伟', 'Alexander Montgomery'])
        self.assertEqual(parsed['sourceFiles'], ['names.docx'])
        self.assertTrue(any('未配置 NAMEPLATE_AI_API_KEY' in item for item in parsed['warnings']))
        self.assertEqual(backend.devices_collection.documents, devices_before)

    def test_docx_ai_failure_falls_back_to_native_text(self):
        document = Document()
        document.add_paragraph('张三')
        document.add_paragraph('李四')
        raw_docx = serialize_test_docx(document)

        with (
            patch.object(backend, 'get_nameplate_ai_api_key', return_value='test-key'),
            patch.object(
                backend, 'call_openai_nameplate_parser',
                side_effect=RuntimeError('upstream unavailable'),
            ),
        ):
            response = self.client.post('/api/nameplates/parse', data={
                'files': (io.BytesIO(raw_docx), 'fallback.docx'),
            }, content_type='multipart/form-data')

        self.assertEqual(response.status_code, 200)
        parsed = response.get_json()['parsed']
        self.assertFalse(parsed['aiUsed'])
        self.assertEqual(parsed['names'], ['张三', '李四'])
        self.assertTrue(any('AI解析失败' in item for item in parsed['warnings']))

    def test_docx_parse_sends_native_text_and_embedded_image_to_ai(self):
        document = Document()
        document.add_paragraph('张三')
        image = Image.new('RGB', (20, 20), (255, 255, 255))
        image_buffer = io.BytesIO()
        image.save(image_buffer, format='PNG')
        image_buffer.seek(0)
        document.add_picture(image_buffer)
        raw_docx = serialize_test_docx(document)

        with (
            patch.object(backend, 'get_nameplate_ai_api_key', return_value='test-key'),
            patch.object(
                backend,
                'call_openai_nameplate_parser',
                return_value=backend.build_nameplate_parse_result(
                    ['张三'], {}, ai_used=True, source_summary='Word AI解析',
                ),
            ) as parser_mock,
        ):
            response = self.client.post('/api/nameplates/parse', data={
                'files': (io.BytesIO(raw_docx), 'mixed.docx'),
                'templateConfig': '{}',
            }, content_type='multipart/form-data')

        self.assertEqual(response.status_code, 200)
        source_text, image_parts, _ = parser_mock.call_args.args
        self.assertIn('张三', source_text)
        self.assertEqual(len(image_parts), 1)
        self.assertEqual(image_parts[0]['mimeType'], 'image/png')
        self.assertTrue(response.get_json()['parsed']['aiUsed'])

    def test_invalid_docx_and_legacy_doc_return_actionable_errors(self):
        invalid_response = self.client.post('/api/nameplates/parse', data={
            'files': (io.BytesIO(b'not-a-docx'), 'broken.docx'),
        }, content_type='multipart/form-data')
        self.assertEqual(invalid_response.status_code, 400)
        self.assertIn('有效的 DOCX', invalid_response.get_json()['error'])

        legacy_response = self.client.post('/api/nameplates/parse', data={
            'files': (io.BytesIO(b'legacy-doc'), 'legacy.doc'),
        }, content_type='multipart/form-data')
        self.assertEqual(legacy_response.status_code, 400)
        self.assertIn('另存为 .docx', legacy_response.get_json()['error'])

    def test_docx_rejects_expansion_entry_count_and_unsafe_paths(self):
        document = Document()
        document.add_paragraph('张三')
        raw_docx = serialize_test_docx(document)

        with patch.object(backend, 'NAMEPLATE_MAX_DOCX_UNCOMPRESSED_BYTES', 128):
            with self.assertRaisesRegex(backend.NameplateParseInputError, '解压后体积'):
                backend.validate_docx_archive(raw_docx)
        with patch.object(backend, 'NAMEPLATE_MAX_DOCX_ENTRIES', 1):
            with self.assertRaisesRegex(backend.NameplateParseInputError, '文件项过多'):
                backend.validate_docx_archive(raw_docx)

        archive_buffer = io.BytesIO()
        with zipfile.ZipFile(archive_buffer, 'w') as archive:
            archive.writestr('[Content_Types].xml', '<Types/>')
            archive.writestr('word/document.xml', '<document/>')
            archive.writestr('../escape.txt', 'blocked')
        with self.assertRaisesRegex(backend.NameplateParseInputError, '异常文件路径'):
            backend.validate_docx_archive(archive_buffer.getvalue())

    def test_sleep_interval_is_clamped(self):
        self.assertEqual(
            backend.normalize_sleep_interval(1, backend.DEFAULT_SLEEP_INTERVAL_SECONDS),
            backend.MIN_SLEEP_INTERVAL_SECONDS,
        )
        self.assertEqual(
            backend.normalize_sleep_interval(10 ** 12, backend.DEFAULT_SLEEP_INTERVAL_SECONDS),
            backend.MAX_SLEEP_INTERVAL_SECONDS,
        )


class NameplateRenderContractTests(unittest.TestCase):
    def test_nameplate_logo_config_validates_image_and_position(self):
        logo_data_url = make_test_logo_data_url()
        config = backend.normalize_nameplate_template_config({
            'logoDataUrl': logo_data_url,
            'logoFileName': '  event-logo.png  ',
            'logoX': -10,
            'logoY': 999,
            'companyX': 999,
            'companyPositionMode': 'custom',
        })

        self.assertEqual(config['logoDataUrl'], logo_data_url)
        self.assertEqual(config['logoFileName'], 'event-logo.png')
        self.assertEqual(config['logoX'], 0)
        self.assertEqual(config['logoY'], 479)
        self.assertEqual(config['companyX'], 799)
        self.assertEqual(config['companyPositionMode'], 'custom')
        self.assertNotIn(
            'companyX',
            backend.normalize_nameplate_template_config({'companyX': 123}),
        )
        self.assertNotIn(
            'logoDataUrl',
            backend.normalize_nameplate_template_config({
                'logoDataUrl': 'data:image/svg+xml;base64,PHN2Zz48L3N2Zz4=',
            }),
        )

    def test_custom_nameplate_logo_uses_saved_canvas_position(self):
        logo_data_url = make_test_logo_data_url()
        image = template_renderer.render_template_image('nameplate', {
            'name': '张三',
            'backgroundStyle': 'formal_red',
            'logoDataUrl': logo_data_url,
            'logoX': 10,
            'logoY': 20,
        })

        self.assertEqual(image.getpixel((100, 30)), (0, 0, 255))

    def test_nameplate_company_uses_saved_horizontal_position_and_stays_visible(self):
        self.assertEqual(
            template_renderer._resolve_nameplate_company_x(
                {'companyX': 24, 'companyPositionMode': 'custom'}, 120,
            ),
            24,
        )
        self.assertEqual(
            template_renderer._resolve_nameplate_company_x(
                {'companyX': 999, 'companyPositionMode': 'custom'}, 120,
            ),
            680,
        )
        self.assertEqual(
            template_renderer._resolve_nameplate_company_x({}, 120),
            340,
        )
        self.assertEqual(
            template_renderer._resolve_nameplate_company_x({'companyX': 24}, 120),
            340,
        )
        self.assertEqual(
            template_renderer._resolve_nameplate_company_x(
                {'companyX': 260, 'companyPositionMode': 'custom'},
                180,
                reference_width=390,
            ),
            365,
        )

        default_image = template_renderer.render_template_image('nameplate', {
            'name': '张三',
            'backgroundStyle': 'formal_red',
            'subtitle': 'ACME',
        })
        moved_image = template_renderer.render_template_image('nameplate', {
            'name': '张三',
            'backgroundStyle': 'formal_red',
            'subtitle': 'ACME',
            'companyX': 650,
            'companyPositionMode': 'custom',
        })
        self.assertNotEqual(default_image.tobytes(), moved_image.tobytes())

    def test_real_nameplate_render_matches_firmware_contract(self):
        result = backend.render_template_with_preview('nameplate', {
            'name': '张三',
            'backgroundStyle': 'formal_red',
            'title': '技术专家',
            'subtitle': '现象光伏',
            'sleepIntervalSeconds': 43200,
        })
        epd_data = result.get('epdData')

        self.assertEqual(len(epd_data), backend.EPD_EXPECTED_CHARS)
        self.assertLessEqual(set(epd_data), backend.EPD_ALLOWED_CHARS)
        self.assertTrue(result.get('previewImage'))


class IndexMigrationTests(unittest.TestCase):
    def test_legacy_global_template_index_is_replaced(self):
        database = FakeDatabase()
        collection = database['saved_nameplate_templates']
        collection.indexes.append({
            'name': 'templateId_1', 'key': {'templateId': 1}, 'unique': True
        })

        db_indexes.ensure_all_indexes(database)

        names = {index['name'] for index in collection.indexes}
        self.assertNotIn('templateId_1', names)
        self.assertIn('owner_template_base_unique', names)
        status_index_names = {
            index['name'] for index in database['device_status'].indexes
        }
        self.assertIn('unclaimedExpiresAt_ttl', status_index_names)


class ConfigValidationTests(unittest.TestCase):
    def test_public_base_url_is_normalized_to_an_origin(self):
        self.assertEqual(
            backend_config._validated_public_base_url('https://Example.COM:8443/?'),
            'https://example.com:8443',
        )

    def test_public_base_url_rejects_invalid_port(self):
        with self.assertRaises(ValueError):
            backend_config._validated_public_base_url('https://example.com:not-a-port')


if __name__ == '__main__':
    unittest.main()
